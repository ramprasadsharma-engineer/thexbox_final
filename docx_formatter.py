"""Reformat and redesign a .docx file into a professional document.

This script uses python-docx to detect document type (book, poem, or notes)
and apply professional formatting suitable for that type. It prompts the user
for an input .docx path and saves a reformatted copy with `_professional` in the
filename.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt, Mm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def detect_document_type(doc: Document) -> str:
    """Infer document type based on simple heuristics.

    Returns "book", "poem", or "notes".
    """
    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    num_paras = len(paragraphs)
    bullet_like = sum(1 for p in paragraphs if p.style.name.startswith("List"))
    avg_len = (
        sum(len(p.text.split()) for p in paragraphs) / num_paras if num_paras else 0
    )
    empty_paras = len(doc.paragraphs) - num_paras

    if bullet_like > num_paras * 0.2:
        return "notes"
    if avg_len < 8 and empty_paras > len(doc.paragraphs) * 0.3:
        return "poem"
    return "book"


def _set_base_formatting(doc: Document) -> None:
    """Apply page size, margins, and default fonts."""
    for section in doc.sections:
        section.page_width = Mm(148)
        section.page_height = Mm(210)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for p in doc.paragraphs:
        for run in p.runs:
            run.font.name = "Times New Roman"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
            run.font.size = Pt(12)


def _format_headings(doc: Document) -> None:
    """Set headings to a sans-serif font."""
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            for run in p.runs:
                run.font.name = "Arial"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                run.font.size = Pt(14)


def _add_page_numbers(doc: Document) -> None:
    """Add centered page numbers to the footer."""
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = paragraph.add_run()
        fld_char1 = OxmlElement("w:fldChar")
        fld_char1.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = "PAGE"
        fld_char2 = OxmlElement("w:fldChar")
        fld_char2.set(qn("w:fldCharType"), "end")
        run._r.extend([fld_char1, instr_text, fld_char2])


def apply_professional_formatting(doc: Document, doc_type: str) -> Document:
    """Apply formatting based on detected document type."""
    _set_base_formatting(doc)

    if doc_type == "book":
        # Assume first paragraph is title; style as Title
        if doc.paragraphs:
            title = doc.paragraphs[0]
            title.style = doc.styles["Title"]
        # Format chapter headings
        for p in doc.paragraphs:
            if re.match(r"^(Chapter|CHAPTER)\b", p.text.strip()):
                p.style = doc.styles["Heading 1"]
    elif doc_type == "poem":
        for p in doc.paragraphs:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif doc_type == "notes":
        for p in doc.paragraphs:
            text = p.text.lstrip()
            if re.match(r"[-*] \s*", text):
                p.style = doc.styles["List Bullet"]
                p.text = re.sub(r"^[-*] \s*", "", text)
            elif re.match(r"\d+[.)] \s*", text):
                p.style = doc.styles["List Number"]
                p.text = re.sub(r"^\d+[.)] \s*", "", text)

    _format_headings(doc)
    _add_page_numbers(doc)
    return doc


def main() -> None:
    """Prompt for a .docx file, reformat it, and save a new version."""
    path_str = input("Enter the path of the .docx file: ")
    path = Path(path_str).expanduser().resolve()
    if not path.exists() or path.suffix.lower() != ".docx":
        print("File not found or not a .docx file.")
        return

    doc = Document(str(path))
    doc_type = detect_document_type(doc)
    new_doc = apply_professional_formatting(doc, doc_type)

    new_path = path.with_stem(path.stem + "_professional")
    new_doc.save(str(new_path))
    print(f"File successfully reformatted. New file saved at: {new_path}")


if __name__ == "__main__":
    main()
